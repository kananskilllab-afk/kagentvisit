"""Generates MERIDIAN_MILESTONE_PLAN.docx — the phase-by-phase milestone document
for the Meridian UI rollout + Action Item Tracker + Agent History features."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


NAVY = RGBColor(0x1E, 0x1B, 0x4B)
BLUE = RGBColor(0x3B, 0x82, 0xF6)
GOLD = RGBColor(0xF5, 0x9E, 0x0B)
GREEN = RGBColor(0x10, 0xB9, 0x81)
RED = RGBColor(0xEF, 0x44, 0x44)
TEXT = RGBColor(0x0F, 0x17, 0x2A)
SUB = RGBColor(0x64, 0x74, 0x8B)
MUTED = RGBColor(0x94, 0xA3, 0xB8)


def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_heading(doc, text, level=1, color=NAVY):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = color
    run.font.name = "Calibri"
    if level == 0:
        run.font.size = Pt(28)
    elif level == 1:
        run.font.size = Pt(20)
    elif level == 2:
        run.font.size = Pt(15)
    elif level == 3:
        run.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_para(doc, text, bold=False, color=TEXT, size=10.5, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(4)
    return p


def add_bullet(doc, text, level=0, bold_lead=None, color=TEXT):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    if bold_lead:
        r1 = p.add_run(bold_lead)
        r1.bold = True
        r1.font.size = Pt(10.5)
        r1.font.color.rgb = NAVY
        r1.font.name = "Calibri"
        r2 = p.add_run(" " + text)
        r2.font.size = Pt(10.5)
        r2.font.color.rgb = color
        r2.font.name = "Calibri"
    else:
        run = p.add_run(text)
        run.font.size = Pt(10.5)
        run.font.color.rgb = color
        run.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(2)
    return p


def add_kv_table(doc, rows, col_widths=(Inches(1.6), Inches(4.6))):
    t = doc.add_table(rows=len(rows), cols=2)
    t.autofit = False
    for i, (k, v) in enumerate(rows):
        c0 = t.cell(i, 0)
        c1 = t.cell(i, 1)
        c0.width = col_widths[0]
        c1.width = col_widths[1]
        set_cell_bg(c0, "F1F5F9")
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(k)
        r0.bold = True
        r0.font.size = Pt(9.5)
        r0.font.color.rgb = NAVY
        r0.font.name = "Calibri"
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(v)
        r1.font.size = Pt(10)
        r1.font.color.rgb = TEXT
        r1.font.name = "Calibri"
    doc.add_paragraph()


def add_phase_header(doc, code, title, duration, deps):
    add_heading(doc, f"{code}  ·  {title}", level=2, color=NAVY)
    add_kv_table(doc, [
        ("Duration", duration),
        ("Depends on", deps),
    ])


def add_section_title(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = BLUE
    run.font.name = "Calibri"
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "3B82F6")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)


def main():
    doc = Document()

    section = doc.sections[0]
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    # ─── Cover ──────────────────────────────────────────────
    cover_p = doc.add_paragraph()
    cover_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = cover_p.add_run("KANAN AVS")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = GOLD
    r.font.name = "Calibri"

    add_heading(doc, "Meridian UI + Action Item Tracker + Agent History", level=0, color=NAVY)
    add_para(doc, "Milestone Implementation Plan", color=SUB, size=14)
    add_para(doc, "Phases P1 → P9  ·  Generated 2026-05-01", color=MUTED, size=10, italic=True)

    add_kv_table(doc, [
        ("Owner", "Omzala"),
        ("Status", "Plan approved — pending /gsd:new-milestone kickoff"),
        ("Total phases", "9"),
        ("Estimated effort", "5–8 working days (sequential, single dev)"),
        ("Branch strategy", "One PR per phase, feature-flagged where possible"),
        ("Source design", "newui/AVS Meridian.html + meridian/{shared,layout,screens1-3}.jsx"),
    ])

    # ─── Confirmed answers ──────────────────────────────────
    add_heading(doc, "Confirmed Decisions", level=1)
    add_section_title(doc, "Action Item Tracker")
    add_bullet(doc, "{ text, assignee, dueDate, status, history[] }", bold_lead="Item shape:")
    add_bullet(doc, "Visit owner, manager/HOD reviewing the visit, and anyone on the agent's account team", bold_lead="Who can mark done:")
    add_bullet(doc, "Append-only history (audit log) — every status change and note edit is recorded", bold_lead="Note model:")
    add_bullet(doc, "Yes — overdue items surface on dashboard widget AND fire in-app notification", bold_lead="Overdue surfacing:")
    add_bullet(doc, "Open items auto-surface on next visit to same agent (\"Open from last visit\" pre-fill block) — orchestrator's call, accepted", bold_lead="Carry-forward:")

    add_section_title(doc, "Agent History")
    add_bullet(doc, "Last 5 visits inline + \"View all\" button → modal/drawer with full timeline", bold_lead="Window:")
    add_bullet(doc, "Open action items count = pending follow-ups (no extra field)", bold_lead="Follow-up definition:")
    add_bullet(doc, "Hidden from home_visit role; visible to superadmin / admin / user / hod / regional_bdm / accounts (read-only for accounts)", bold_lead="Role visibility:")

    add_section_title(doc, "Meridian UI Rollout")
    add_bullet(doc, "Tailwind — extract Meridian tokens into tailwind.config.js (colors, font-family, shadow, border)", bold_lead="Styling:")
    add_bullet(doc, "Inter (300–800) + Manrope (500–900) via Google Fonts in index.html, latest available", bold_lead="Fonts:")
    add_bullet(doc, "Keep the existing logo (do NOT swap to newui/logo.png)", bold_lead="Logo:")
    add_bullet(doc, "Yes — every module: DailyReport, PostDemoFeedback, PostFieldDay, PostInPersonVisit, Profile, FormsHub, FormsAdmin included", bold_lead="Coverage:")

    # ─── Architecture overview ──────────────────────────────
    add_heading(doc, "Architecture Overview", level=1)
    add_section_title(doc, "Backend changes")
    add_bullet(doc, "VisitSchedule.actionItems[] subdocument added (text, assignee, dueDate, status, history[]); Visit model gets the same if standalone visits also need it")
    add_bullet(doc, "FormConfig: new field type 'action_items' so the existing dynamic-form engine can render the tracker at any step the admin places it (currently described as 'step 10' but config-driven)")
    add_bullet(doc, "New endpoints: PATCH /api/visits/:visitId/action-items/:itemId, POST /api/visits/:visitId/action-items, GET /api/agents/:agentId/history")
    add_bullet(doc, "Notification model: new type 'action_item_overdue' (daily cron job to enqueue)")
    add_bullet(doc, "AuditLog: new targetModel 'ActionItem'")

    add_section_title(doc, "Frontend changes")
    add_bullet(doc, "Design-system module at client/src/design/ (tokens.js, primitives: Card, Btn, Lbl, Input, Avatar, StatusBadge, EmptyState, Sidebar, TopBar, AppShell)")
    add_bullet(doc, "Tailwind config extended with meridian palette, font families, custom shadows")
    add_bullet(doc, "Every page rebuilt to consume the new primitives — Tailwind classes, no inline-style copy from the mock")
    add_bullet(doc, "PlanModal agent step gets <AgentHistoryCard agentId={...} /> embedded under each agent")
    add_bullet(doc, "VisitsList detail panel gets <ActionItemTracker visitId={...} /> tab")
    add_bullet(doc, "Dashboard gets \"My Open Action Items\" widget (overdue count + list) for all roles except accounts")

    # ─── Phase pages ────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "Phase Breakdown", level=1)

    # P1
    add_phase_header(doc, "P1", "Design-System Extraction & Foundation", "0.75 day", "—")
    add_section_title(doc, "Goal")
    add_para(doc, "Establish the Meridian design language as a reusable, Tailwind-native primitive set so every following phase consumes the same building blocks.")
    add_section_title(doc, "Backend")
    add_bullet(doc, "(none)")
    add_section_title(doc, "Frontend")
    add_bullet(doc, "Add Inter + Manrope to client/index.html via Google Fonts")
    add_bullet(doc, "Extend tailwind.config.js: colors.meridian.{navy, blue, gold, green, red, sky, purple, bg, border, text, sub, muted, rowHov}; fontFamily.display = ['Manrope', ...]; boxShadow.card; borderRadius scale match")
    add_bullet(doc, "client/src/design/tokens.js exports the same constants for runtime use (charts, dynamic styles)")
    add_bullet(doc, "client/src/design/icons.jsx — port the IC{} catalog from newui/meridian/shared.jsx as named exports plus the <Icon> primitive")
    add_bullet(doc, "client/src/design/primitives/ — Card.jsx, Btn.jsx, Input.jsx, Lbl.jsx, Avatar.jsx, StatusBadge.jsx, EmptyState.jsx, NotifBell.jsx, RolePill.jsx, SectionTitle.jsx, SparkArea.jsx, SparkBar.jsx")
    add_bullet(doc, "client/src/design/index.js barrel export")
    add_bullet(doc, "Storybook-style preview page at /design-system route (dev only) so we can eyeball every primitive in isolation")
    add_section_title(doc, "Acceptance")
    add_bullet(doc, "All primitives render at /design-system and visually match the newui mock")
    add_bullet(doc, "Existing pages still work unchanged (no consumer migrations yet)")
    add_bullet(doc, "tailwind class meridian-navy applies #1E1B4B")

    # P2
    add_phase_header(doc, "P2", "Action Item Tracker — Backend & Data Model", "1 day", "P1")
    add_section_title(doc, "Goal")
    add_para(doc, "Persist action items per visit with full audit history. Expose REST endpoints. No UI yet — Postman/curl verifiable.")
    add_section_title(doc, "Backend")
    add_bullet(doc, "VisitSchedule.js: add actionItems: [{ _id, text, assignee: ObjectId(User), dueDate, status: 'pending'|'done'|'cancelled', createdBy, createdAt, history: [{ at, by, action, fromStatus, toStatus, note }] }]")
    add_bullet(doc, "Visit.js: same actionItems[] block (mirrors VisitSchedule for legacy standalone visits)")
    add_bullet(doc, "Helper services/actionItem.service.js: addItem, updateItem, changeStatus, listOpenForAgent, listOverdue (each appends to history[])")
    add_bullet(doc, "Routes: routes/actionItems.routes.js mounted at /api/visits/:visitId/action-items (POST, GET, PATCH /:id, DELETE /:id, POST /:id/status)")
    add_bullet(doc, "Permissions middleware: visit owner OR (manager/hod with org match) OR account-team membership of agent")
    add_bullet(doc, "AuditLog targetModel enum: add 'ActionItem'")
    add_bullet(doc, "Notification model: new kind 'action_item_overdue' with payload { visitId, itemId, agentName, dueDate, daysOverdue }")
    add_bullet(doc, "Cron at server/jobs/actionItemOverdue.job.js — runs daily 09:00 server time, enqueues notifications for items where status=pending && dueDate < now")
    add_bullet(doc, "Migration script server/migration/seedActionItemsField.js (no-op for existing data; ensures index on actionItems.status + actionItems.dueDate)")
    add_section_title(doc, "Frontend")
    add_bullet(doc, "(none in this phase — backend only)")
    add_section_title(doc, "Acceptance")
    add_bullet(doc, "Postman collection (committed to docs/postman/) covers all CRUD + status-change paths")
    add_bullet(doc, "AuditLog rows appear for every mutation")
    add_bullet(doc, "Cron dry-run produces correct overdue list against a fixture set")

    # P3
    add_phase_header(doc, "P3", "Action Item Tracker — Frontend & FormConfig Field Type", "1 day", "P1, P2")
    add_section_title(doc, "Goal")
    add_para(doc, "Render the tracker inside NewVisit (configurable via FormBuilder) and inside the VisitsList detail panel. Done/note workflows fully usable.")
    add_section_title(doc, "Backend")
    add_bullet(doc, "FormConfig.js: extend field.type enum with 'action_items'")
    add_bullet(doc, "FormBuilder admin endpoint accepts the new type")
    add_section_title(doc, "Frontend")
    add_bullet(doc, "client/src/components/SurveyForm/fields/ActionItemsField.jsx — add-one-by-one input with enter-to-add, inline edit, drag-reorder optional, due-date picker, assignee select (loads org users)")
    add_bullet(doc, "NewVisit form renderer wires the new field type to ActionItemsField")
    add_bullet(doc, "client/src/components/ActionItemTracker.jsx — used in VisitsList detail panel and VisitDetailModal: list with status toggle, write-note dialog (append-only), history disclosure")
    add_bullet(doc, "VisitsList.jsx detail drawer gains an \"Action Items\" tab next to existing tabs")
    add_bullet(doc, "FormBuilder.jsx UI palette: drag the new \"Action Items\" field block into any step")
    add_bullet(doc, "Dashboard widget: \"My Open Action Items\" card (top-3 + count link)")
    add_section_title(doc, "Acceptance")
    add_bullet(doc, "Admin can add the action_items field to step 10 (or any step) via FormBuilder — DB-driven, not hardcoded")
    add_bullet(doc, "Field agent fills items at submit, manager/HOD marks done with note from VisitsList, history shows the trail")
    add_bullet(doc, "Overdue notification appears in NotifBell within 1 minute of cron firing (e2e test)")

    # P4
    add_phase_header(doc, "P4", "Agent History — Backend Aggregation & API", "0.5 day", "P2")
    add_section_title(doc, "Goal")
    add_para(doc, "Single endpoint returns everything PlanModal and ManageAgent need to show \"what happened with this agent so far.\"")
    add_section_title(doc, "Backend")
    add_bullet(doc, "controllers/agents.controller.js: getAgentHistory(agentId) → { agent, lastVisitAt, totalVisits, openActionItemsCount, openActionItems[], recentVisits[5], allVisitsCount, lastBdm, avgRating, plansCount }")
    add_bullet(doc, "Mongo aggregation pipeline against Visit + VisitSchedule + VisitPlan with $facet for the counts; lean+projection to keep payload small")
    add_bullet(doc, "Route: GET /api/agents/:agentId/history (all roles except home_visit; accounts read-only)")
    add_bullet(doc, "Route: GET /api/agents/:agentId/visits?limit=N for the \"View all\" expansion")
    add_bullet(doc, "Cache layer: in-memory 60s TTL keyed by agentId (small-team app, fine)")
    add_section_title(doc, "Frontend")
    add_bullet(doc, "(none — wired in P5)")
    add_section_title(doc, "Acceptance")
    add_bullet(doc, "Endpoint returns under 200ms for an agent with 50 visits")
    add_bullet(doc, "home_visit role gets 403; accounts gets 200 read-only")

    # P5
    add_phase_header(doc, "P5", "Agent History — UI in PlanModal & ManageAgent", "0.75 day", "P1, P4")
    add_section_title(doc, "Goal")
    add_para(doc, "Surface the history wherever an agent is selected so the user knows context before scheduling.")
    add_section_title(doc, "Backend")
    add_bullet(doc, "(none)")
    add_section_title(doc, "Frontend")
    add_bullet(doc, "client/src/components/AgentHistoryCard.jsx — KPI row (total visits, last visit, open items, avg rating) + last-5 visit list + \"View all\" button → AgentVisitsDrawer")
    add_bullet(doc, "client/src/components/AgentVisitsDrawer.jsx — paginated full history with status filter")
    add_bullet(doc, "PlanModal.jsx step 2 (Agents): under each selected agent chip, expandable AgentHistoryCard")
    add_bullet(doc, "ManageAgent.jsx detail drawer gains a \"History\" tab using the same component")
    add_bullet(doc, "NewVisit.jsx — when agent is selected on the institution step, show a compact AgentHistoryCard variant inline")
    add_bullet(doc, "Role guard: hide for home_visit users via the existing role context")
    add_section_title(doc, "Acceptance")
    add_bullet(doc, "Selecting an agent in PlanModal shows last-5 visits, open action items count, and last visit date within 300ms")
    add_bullet(doc, "\"View all\" opens drawer with paginated full history")
    add_bullet(doc, "home_visit user sees PlanModal without the card section (no layout gaps)")

    # P6
    add_phase_header(doc, "P6", "Layout + Login + Dashboard — Meridian Rebuild", "1 day", "P1, P3 (for widget)")
    add_section_title(doc, "Goal")
    add_para(doc, "Visible Meridian transformation begins. Shell, login, and home land first because they frame everything.")
    add_section_title(doc, "Backend")
    add_bullet(doc, "(none)")
    add_section_title(doc, "Frontend")
    add_bullet(doc, "Layout.jsx replaced by Meridian Sidebar (236px navy, gold active accent) + TopBar (58px) — using P1 primitives, Tailwind only")
    add_bullet(doc, "Login.jsx: split-pane layout — left navy panel with stats, right form with Inter inputs, gold underline accent")
    add_bullet(doc, "Dashboard.jsx: 4 KPI cards with left-accent borders + 2-up charts row (visit activity SparkBar + status breakdown) + recent visits table + \"My Open Action Items\" widget")
    add_bullet(doc, "Page transition animation (slideIn 0.22s) wired into Layout outlet")
    add_bullet(doc, "Mobile: sidebar collapses to icon-only at <1024px, drawer at <768px")
    add_section_title(doc, "Acceptance")
    add_bullet(doc, "Visual side-by-side review against newui/AVS Meridian.html — matches within token tolerance")
    add_bullet(doc, "All existing routes still resolve")
    add_bullet(doc, "Manual smoke: login → dashboard → click each sidebar item → no broken pages (other pages may still be old-style — fine, fixed in later phases)")

    # P7
    add_phase_header(doc, "P7", "Visits Module — VisitsList + NewVisit + VisitDetailModal", "1 day", "P1, P3, P5")
    add_section_title(doc, "Goal")
    add_para(doc, "The most-used surface gets the Meridian treatment, with the new Action Items tab and inline Agent History.")
    add_section_title(doc, "Backend")
    add_bullet(doc, "(none)")
    add_section_title(doc, "Frontend")
    add_bullet(doc, "VisitsList.jsx — Meridian table (zebra hover, status pills, avatar institution column, rating stars, side detail panel pattern)")
    add_bullet(doc, "Filter bar: search input with icon + status chip group + Export/New Visit primary buttons")
    add_bullet(doc, "Detail drawer tabs: Summary / Action Items / Expenses / Documents")
    add_bullet(doc, "NewVisit.jsx — chrome wrapped in Meridian shell, StepIndicator restyled (gold dot for current, green for done), card body uses meridian.card shadow")
    add_bullet(doc, "VisitDetailModal.jsx — full-screen Meridian modal with sticky header")
    add_bullet(doc, "Visit history bar gets the new \"Action Item Tracker\" entry from P3")
    add_section_title(doc, "Acceptance")
    add_bullet(doc, "End-to-end: create plan → agent history shows → fill visit → action items captured → submit → review in VisitsList → mark items done")

    # P8
    add_phase_header(doc, "P8", "Calendar + PlanModal + VisitPlanDetail", "0.75 day", "P1, P5")
    add_section_title(doc, "Goal")
    add_para(doc, "Planning surfaces share the same look. PlanModal already gets AgentHistoryCard from P5; this phase finishes the chrome.")
    add_section_title(doc, "Backend")
    add_bullet(doc, "(none)")
    add_section_title(doc, "Frontend")
    add_bullet(doc, "Calendar.jsx — Month/Week/Day/Agenda views restyled with Meridian event chips (color-by-status), agenda list with avatar+pill rows, top toolbar uses Btn variants")
    add_bullet(doc, "PlanModal.jsx — 4-step stepper restyled (numbered nav with gold current, navy completed)")
    add_bullet(doc, "VisitPlanDetail.jsx — overview KPI cards, balance bar, tabs in Meridian style; Documents tab uses Card + grid")
    add_bullet(doc, "ScheduleModal.jsx — modal restyled to match Meridian dialog pattern")
    add_section_title(doc, "Acceptance")
    add_bullet(doc, "Calendar views render correctly across viewports; keyboard shortcuts preserved")

    # P9
    add_phase_header(doc, "P9", "Expenses + Claims + ClaimDetail", "0.75 day", "P1")
    add_section_title(doc, "Goal")
    add_para(doc, "Money screens get the treatment with strong status pills (paid/approved/needs_justification).")
    add_section_title(doc, "Backend")
    add_bullet(doc, "(none)")
    add_section_title(doc, "Frontend")
    add_bullet(doc, "Expenses/* pages — list, NewClaim wizard, ClaimDetail; category icons styled per CAT_META")
    add_bullet(doc, "Policy violation banner restyled (red Meridian alert pattern)")
    add_bullet(doc, "Templates picker uses Meridian Card grid")
    add_bullet(doc, "Receipt upload zone restyled")
    add_section_title(doc, "Acceptance")
    add_bullet(doc, "Submit flow still triggers policy.evaluateClaim() and surfaces violations correctly")

    # P10 (renumbered as continuation under P9 heading or extra)
    add_phase_header(doc, "P10", "ManageAgent + Users + FormBuilder + Analytics", "0.75 day", "P1, P5")
    add_section_title(doc, "Goal")
    add_para(doc, "Admin surfaces. Users + FormBuilder are superadmin-only — visual only, no logic changes.")
    add_section_title(doc, "Backend")
    add_bullet(doc, "(none)")
    add_section_title(doc, "Frontend")
    add_bullet(doc, "ManageAgent.jsx — agent table with avatar, rank pill, type pill, active toggle; detail drawer with History tab (P5 component)")
    add_bullet(doc, "SuperAdmin/UserManagement.jsx — Meridian table with role pills (ROLE_META) and active dot")
    add_bullet(doc, "FormBuilder.jsx — left palette / right canvas restyled; the new action_items block sits in the palette")
    add_bullet(doc, "Analytics.jsx — KPI cards + SparkArea + SparkBar; export PDF button")
    add_section_title(doc, "Acceptance")
    add_bullet(doc, "Superadmin can still create/edit/disable users; FormBuilder still saves configs")

    # P11
    add_phase_header(doc, "P11", "Tail Pages + Cleanup + Regression Sweep", "0.5 day", "P1–P10")
    add_section_title(doc, "Goal")
    add_para(doc, "Catch-all for the screens not in the newui mockups, then a global sweep for visual regressions.")
    add_section_title(doc, "Backend")
    add_bullet(doc, "(none)")
    add_section_title(doc, "Frontend")
    add_bullet(doc, "DailyReport.jsx, PostDemoFeedback.jsx, PostFieldDay.jsx, PostInPersonVisit.jsx, Profile.jsx, FormsHub.jsx, FormsAdmin.jsx — apply Meridian primitives consistently (mostly form Cards + Btn + Lbl)")
    add_bullet(doc, "Remove dead Tailwind classes / unused custom CSS from index.css")
    add_bullet(doc, "Visual regression pass: every page screenshot side-by-side against the mock")
    add_bullet(doc, "Accessibility sweep: focus rings, aria-labels on icon buttons, contrast check on gold-on-navy")
    add_section_title(doc, "Acceptance")
    add_bullet(doc, "No page in the app still uses the pre-Meridian look")
    add_bullet(doc, "Lighthouse a11y score >= 90 on dashboard, visits, expenses")

    # ─── Risk register ──────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "Risk Register & Mitigations", level=1)

    risks = [
        ("FormConfig migration risk", "High",
         "Adding a new field type 'action_items' could break form rendering for existing visits that don't have it.",
         "Default the field to optional; renderer falls back to null-safe; migration script is idempotent. Test with a copy of prod FormConfig before merge."),
        ("Action item history bloat", "Med",
         "Append-only history can grow unbounded for hot agents.",
         "Cap displayed history at 20 entries with \"Show all\" expansion; archive entries older than 1 year via monthly cron."),
        ("Tailwind config conflict", "Med",
         "Adding meridian palette may shadow existing brand-blue / brand-navy class names.",
         "Namespace under `meridian.*` in config, migrate consumers explicitly, keep a compat shim alias for one phase."),
        ("Permission drift", "Med",
         "New endpoints (action items, agent history) need consistent role checks.",
         "Centralize in a single requireRole() middleware; unit-test all 6 roles × 4 endpoints matrix."),
        ("Notification spam", "Low",
         "Daily overdue cron could fire dozens of notifications for the same user.",
         "Coalesce into one digest notification per user per day; deep-link to filtered Action Items view."),
        ("UI rebuild regressions", "Med",
         "Tailwind class changes can break responsive behavior on tablets.",
         "Test at 1440 / 1024 / 768 / 414 widths per page; phase ships behind ?meridian=1 query flag during P6–P11 for quick rollback."),
        ("Inline-style → Tailwind translation drift", "Low",
         "Numeric values in the mock (height:36) need Tailwind equivalents.",
         "P1 produces a tokens reference table mapping every numeric value to a Tailwind class; reused by every subsequent phase."),
    ]
    t = doc.add_table(rows=1 + len(risks), cols=4)
    t.autofit = False
    headers = ["Risk", "Severity", "Impact", "Mitigation"]
    widths = [Inches(1.4), Inches(0.7), Inches(2.0), Inches(2.4)]
    for i, h in enumerate(headers):
        c = t.cell(0, i)
        c.width = widths[i]
        set_cell_bg(c, "1E1B4B")
        p = c.paragraphs[0]
        r = p.add_run(h.upper())
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(9)
        r.font.name = "Calibri"
    for ri, row in enumerate(risks, 1):
        for ci, val in enumerate(row):
            c = t.cell(ri, ci)
            c.width = widths[ci]
            p = c.paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(9.5)
            r.font.color.rgb = TEXT
            r.font.name = "Calibri"
            if ci == 1:
                r.bold = True
                if val == "High":
                    r.font.color.rgb = RED
                elif val == "Med":
                    r.font.color.rgb = GOLD
                else:
                    r.font.color.rgb = GREEN

    # ─── File-touch matrix ──────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "File-Touch Matrix (preview)", level=1)
    add_para(doc, "Authoritative list will be regenerated by /gsd:plan-phase per phase. This is the predicted scope.", color=SUB, italic=True)

    matrix = [
        ("P1", "client/index.html, client/tailwind.config.js, client/src/design/* (new), client/src/App.jsx (route /design-system)"),
        ("P2", "server/models/{Visit,VisitSchedule,AuditLog,Notification}.js, server/services/actionItem.service.js (new), server/routes/actionItems.routes.js (new), server/server.js, server/jobs/actionItemOverdue.job.js (new), server/migration/seedActionItemsField.js (new), docs/postman/action-items.json (new)"),
        ("P3", "server/models/FormConfig.js, server/controllers/forms.controller.js, client/src/components/SurveyForm/fields/ActionItemsField.jsx (new), client/src/components/ActionItemTracker.jsx (new), client/src/pages/{NewVisit,VisitsList,FormBuilder,Dashboard}.jsx, client/src/components/VisitDetailModal.jsx"),
        ("P4", "server/controllers/agents.controller.js, server/routes/agents.routes.js, server/services/agentHistory.service.js (new)"),
        ("P5", "client/src/components/{AgentHistoryCard,AgentVisitsDrawer}.jsx (new), client/src/components/PlanModal.jsx, client/src/pages/{ManageAgent,NewVisit}.jsx"),
        ("P6", "client/src/components/Layout.jsx, client/src/pages/{Login,Dashboard}.jsx, client/src/index.css"),
        ("P7", "client/src/pages/{VisitsList,NewVisit}.jsx, client/src/components/{VisitDetailModal,SurveyForm/StepIndicator}.jsx"),
        ("P8", "client/src/pages/{Calendar,VisitPlanDetail}.jsx, client/src/components/{PlanModal,ScheduleModal}.jsx"),
        ("P9", "client/src/pages/Expenses/* (all)"),
        ("P10", "client/src/pages/{ManageAgent,SuperAdmin/UserManagement,FormBuilder,Analytics}.jsx"),
        ("P11", "client/src/pages/{DailyReport,PostDemoFeedback,PostFieldDay,PostInPersonVisit,Profile,FormsHub,FormsAdmin}.jsx, client/src/index.css"),
    ]
    t = doc.add_table(rows=1 + len(matrix), cols=2)
    t.autofit = False
    for i, h in enumerate(["Phase", "Files touched (predicted)"]):
        c = t.cell(0, i)
        c.width = Inches(0.7) if i == 0 else Inches(5.8)
        set_cell_bg(c, "1E1B4B")
        p = c.paragraphs[0]
        r = p.add_run(h.upper())
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(9)
        r.font.name = "Calibri"
    for ri, (phase, files) in enumerate(matrix, 1):
        for ci, val in enumerate([phase, files]):
            c = t.cell(ri, ci)
            c.width = Inches(0.7) if ci == 0 else Inches(5.8)
            p = c.paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(9)
            r.font.color.rgb = TEXT
            r.font.name = "Calibri"
            if ci == 0:
                r.bold = True
                r.font.color.rgb = NAVY

    # ─── Definition of Done ─────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "Definition of Done (per phase)", level=1)
    for line in [
        "Code merged to main with squashed conventional-commit message",
        "All previous-phase pages still work (visual regression checked)",
        "New endpoints documented in docs/postman/ collection",
        "Manual smoke test performed by Omzala — sign off in PR description",
        "No console errors at any visited route",
        "Lighthouse a11y >= 85 on changed pages",
        "Tailwind classes used (no inline-style soup); design tokens come from tailwind config or src/design/tokens.js",
    ]:
        add_bullet(doc, line)

    add_heading(doc, "Out-of-Scope (explicit)", level=1)
    for line in [
        "Mobile-native app — web responsive only",
        "Dark mode — Meridian is light-mode-only by design",
        "Internationalization — copy stays English",
        "Backend rewrite — schema additions only, no breaking changes",
        "Migrating existing visits to populate actionItems[] retroactively — they default to [] and admins seed forward",
    ]:
        add_bullet(doc, line)

    add_heading(doc, "Next Action", level=1)
    add_para(doc,
             "Run /gsd:new-milestone \"Meridian UI + Action Items + Agent History\". "
             "I'll generate PROJECT.md updates and the per-phase RESEARCH/PLAN docs from this blueprint.",
             bold=True, color=NAVY, size=11)

    out = r"C:\Users\Kanan\OneDrive\Desktop\Projects\kagent_visit\kagentvisit\MERIDIAN_MILESTONE_PLAN.docx"
    doc.save(out)
    print(f"WROTE {out}")


if __name__ == "__main__":
    main()
